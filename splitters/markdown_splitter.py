import logging
import os
from pathlib import Path
from typing import List, Literal
from loguru import logger
import docx
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from langchain_core.documents import Document
from langchain_core.output_parsers.json import JsonOutputParser
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)


HEADER_1 = "Header 1"
HEADER_2 = "Header 2"
HEADERS_TO_SPLIT_ON = [
    ("#", HEADER_1),
    ("##", HEADER_2),
    ("###", HEADER_2),
    ("####", HEADER_2),
]

class MarkdownSplitter:
    def __init__(self,) -> None:
        self._output_parser = JsonOutputParser()
        self._md_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=HEADERS_TO_SPLIT_ON
        )
        self._recurcive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=0,
        )

    def _docx_to_markdown(self, docx_path: Path):
        """
        Преобразует документ .docx в формат Markdown.

        Аргументы:
            docx_path (Path): Путь к .docx файлу.

        Возвращает:
            str: Текст в формате Markdown.
        """
        # Extract file name without extension for the main title
        product_name = docx_path.parent.name
        print('[DOCX PATH]', docx_path)
        markdown_lines = [f"# {product_name}"]

        # Open the .docx document
        doc = docx.Document(docx_path)

        for paragraph in doc.paragraphs:
            # Check if the paragraph is a heading or centered text
            if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                # Treat centered text as an ## header
                markdown_lines.append(f"{paragraph.text}\n")
            elif paragraph.style.name.startswith("Heading"):
                # Convert the heading level to Markdown header
                heading_level = (
                    int(paragraph.style.name.split()[-1]) + 1
                )  # Heading 1 -> ##, Heading 2 -> ###, etc.
                markdown_lines.append(f"{paragraph.text}. ")
            else:
                # Normal paragraph text
                markdown_lines.append(paragraph.text + "\n")

        # Join all lines into a single markdown text
        markdown_text = "".join(markdown_lines)

        # Save the output as a Markdown file
        # markdown_file = f"{file_name}.md"
        # with open(markdown_file, "w") as md_file:
        #     md_file.write(markdown_text)

        # print(f"Markdown file saved as {markdown_file}")
        return markdown_text


    def _markdown_split(self, markdown_text: str) -> List[Document]:
        """
        Разделяет текст в формате Markdown на документы.

        Аргументы:
            markdown_text (str): Текст в формате Markdown.

        Возвращает:
            List[Document]: Список документов.
        """
        return self._md_splitter.split_text(markdown_text)

    def _recursive_split(self, text: str) -> list[Document]:
        """
        Рекурсивно разделяет текст на документы.

        Аргументы:
            text (str): Текст для разделения.

        Возвращает:
            list[Document]: Список документов.
        """
        documents: list[Document] = self._recurcive_splitter.transform_documents(text)
        for doc in documents:
            if HEADER_2 in doc.metadata:
                doc.page_content = (
                    f"{doc.page_content}"
                )
                del doc.metadata[HEADER_2]

        return documents

    def _merge_documents(self, documents: List[Document]) -> List[Document]:
        """
        Объединяет документы, если их общий размер меньше порогового значения.

        Аргументы:
            documents (List[Document]): Список документов для объединения.

        Возвращает:
            List[Document]: Список объединенных документов.
        """
        if len(documents) == 0 or len(documents) == 1:
            return documents

        first_doc = documents[0]
        merged_documents = []
        current_doc = None
        size = 512 - len(first_doc.metadata[HEADER_1])

        # Проверяем каждый документ
        for doc in documents:
            if current_doc is None:
                current_doc = Document(
                    page_content=doc.page_content, metadata=doc.metadata.copy()
                )
            elif len(current_doc.page_content) + len(doc.page_content) <= size:
                current_doc.page_content += f"\n{doc.page_content}"
            else:
                merged_documents.append(current_doc)
                current_doc = Document(
                    page_content=doc.page_content, metadata=doc.metadata.copy()
                )

        if current_doc is not None:
            merged_documents.append(current_doc)

        # Add header to each document
        for doc in merged_documents:
            doc.page_content = (
                f"{doc.page_content}"
            )
            del doc.metadata[HEADER_1]
        merged_documents = [chunk.page_content.replace('# AutoRAG','')  for chunk in merged_documents]
        return merged_documents

    def _create_documents_from_docx(self, docx_path: Path) -> List[Document]:
        """
        Создает документы из .docx файла.

        Аргументы:
            docx_path (Path): Путь к .docx файлу.

        Возвращает:
            List[Document]: Список документов.
        """
        logger.info(f"Loading document: {docx_path}")
        markdown_text = self._docx_to_markdown(docx_path)
        logger.info(f"Markdown text size: {len(markdown_text)}")
        md_split = self._markdown_split(markdown_text)
        logger.info(f"Number of documents after markdown split: {len(md_split)}")
        recursive_split = self._recursive_split(md_split)
        logger.info(
            f"Number of documents after recursive split: {len(recursive_split)}"
        )
        merged_docs = self._merge_documents(recursive_split)
        logger.info(f"Number of documents after merge: {len(merged_docs)}")
        return merged_docs

    def prepare_questions_dataset(self, processed_products: list[str]) -> pd.DataFrame:
        """
        Генерирует вопросы для продуктов и сохраняет их в файл Excel.

        Метод выполняет следующие шаги:
        1. Загружает существующий файл, если он существует.
        2. Для каждого продукта генерирует вопросы на основе документов.
        3. Сохраняет сгенерированные вопросы в DataFrame.
        4. Удаляет дубликаты вопросов.
        5. Сохраняет DataFrame в файл Excel.

        Returns:
            pd.DataFrame: DataFrame, содержащий сгенерированные вопросы.
        """
        logger.info("Генерация вопросов для продуктов.")
        questions = []
        # Save questions as a pandas DataFrame
        output_filepath = os.path.join(self._cfg.output_dir, self._cfg.output_filename)

        # Load existing file if it exists
        if os.path.exists(output_filepath):
            df = pd.read_excel(output_filepath)
            # questions = df.to_dict("records")
            logger.info(f"Загружен существующий файл: {output_filepath}")
            # logger.info(f"Количество запросов в файле: {len(questions)}")
            return df

        for product in self._products:
            if product in processed_products:
                continue
            products_questions = []
            logger.info(f"Генерация вопросов для продукта: {product}")
            for docx_file in self._products[product]["docx_files"]:
                documents = self._create_documents_from_docx(docx_file)
                for doc in documents:
                    products_questions.extend(
                        [
                            {
                                self._cfg.topic_column: product,
                                self._cfg.query_column: question,
                            }
                            for question in self._generate_questions(doc)
                        ]
                    )
            # if limits:
            #     random.shuffle(products_questions)
            #     products_questions = products_questions[: limits[product]]

            questions.extend(products_questions)

            df = pd.DataFrame(questions)
            df.drop_duplicates(subset=self._cfg.duplication_subset, inplace=True)
            logger.info(
                f"Количество запросов по темам: \n{df[self._cfg.topic_column].value_counts().to_string()}"
            )
            df.to_excel(os.path.join(output_filepath), index=False)
            logger.info(f"Запросы сохранены в {self._cfg.output_filename}")

        return df


#splitter = MarkdownSplitter()
#chunks = splitter._create_documents_from_docx(Path('test_rag_data\_credit_card_merged_180924.docx'))
#print(chunks)